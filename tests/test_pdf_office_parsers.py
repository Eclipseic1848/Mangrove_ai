# -*- coding: utf-8 -*-
"""PDF/DOCX 解析器测试（Phase 2 Task 6）。"""
from __future__ import annotations

import io
from pathlib import Path

import pytest

from src.data_prep.artifact_store import ArtifactStore
from src.data_prep.models import RawArtifact
from src.data_prep.document_models import DocumentElement
from src.parsers.pdf import PdfParser
from src.parsers.office import OfficeParser
from src.services.mineru_document import (
    MinerUHealth,
    MinerUPageBlock,
    MinerUParseResult,
    MinerUServiceError,
)


def _make_pdf_bytes(pages: list[str]) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    for text in pages:
        pdf.add_page()
        pdf.set_font("helvetica", size=12)
        pdf.cell(0, 10, text)
    return bytes(pdf.output())


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_pdf_parser_extracts_pages_with_position(tmp_path: Path):
    store = ArtifactStore(root=str(tmp_path))
    data = _make_pdf_bytes(["First page content", "Second page content"])
    art = store.write_raw(
        task_id="task-1", source_id="file-1", data=data,
        uri="upload://data.pdf", media_type="application/pdf", ext="pdf",
    )

    parser = PdfParser()
    records, rejects = parser.parse(art, data)

    assert len(records) == 2
    assert "First page content" in records[0].data["text"]
    assert records[0].meta["position"]["page"] == 1
    assert records[1].meta["position"]["page"] == 2
    assert records[0].meta["parser"] == "pdf"
    assert records[0].meta["artifact_id"] == art.artifact_id
    assert records[0].meta["page_kind"] == "digital"
    assert records[0].meta["route"]["primary_backend"] == "docling"
    assert records[0].data["elements"][0]["bbox"]["coordinate_space"] == "pdf_points"
    DocumentElement.model_validate(records[0].data["elements"][0])
    assert rejects == []


def test_artifact_store_accepts_relative_root_for_immutable_json(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Windows 下相对根目录也必须能安全返回第三方缓存引用。"""
    monkeypatch.chdir(tmp_path)
    store = ArtifactStore(root="downloads")

    raw_ref = store.write_json_if_absent(
        "task-relative",
        "third_party/mineru/result.json",
        {"status": "completed"},
    )

    assert raw_ref == "task-relative/third_party/mineru/result.json"
    assert store.resolve_path(raw_ref).exists()


def test_pdf_parser_emits_deterministic_table_rows(tmp_path: Path) -> None:
    """数字型 PDF 表格直接复用 pdfplumber，不再退化成散乱单词。"""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=10)
    for row in (("name", "amount"), ("Alice", "100"), ("Bob", "200")):
        for cell in row:
            pdf.cell(40, 10, cell, border=1)
        pdf.ln()
    data = bytes(pdf.output())
    store = ArtifactStore(root=str(tmp_path))
    artifact = store.write_raw(
        "task-table",
        "file-1",
        data,
        uri="table.pdf",
        media_type="application/pdf",
        ext="pdf",
    )

    records, rejects = PdfParser(
        artifact_store=store,
        use_remote_ocr=False,
    ).parse(artifact, data)
    table_elements = [
        DocumentElement.model_validate(element)
        for record in records
        for element in record.data.get("elements") or []
        if element.get("element_type") == "table"
    ]

    assert rejects == []
    assert len(table_elements) == 3
    assert table_elements[0].extractor == "pdfplumber"
    assert table_elements[0].metadata["table_columns"] == ["列1", "列2"]
    assert table_elements[1].metadata["table_row"] == {
        "列1": "Alice",
        "列2": "100",
    }
    assert table_elements[1].metadata["location"] == {
        "kind": "pdf_table_row",
        "table": 1,
        "row": 2,
    }


def test_pdf_parser_marks_empty_page_as_ocr_required(tmp_path: Path):
    """无数字文本的页面进入 rejects（ocr_required），不静默消失。"""
    store = ArtifactStore(root=str(tmp_path))
    # 空白页（无文本）+ 有文本页
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()  # 空白页
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    pdf.cell(0, 10, "Has text")
    data = bytes(pdf.output())
    art = store.write_raw(
        task_id="task-1", source_id="file-1", data=data,
        uri="upload://data.pdf", media_type="application/pdf", ext="pdf",
    )

    parser = PdfParser()
    records, rejects = parser.parse(art, data)

    assert len(records) == 1  # 只有第 2 页有文本
    assert records[0].meta["position"]["page"] == 2
    assert len(rejects) == 1
    assert rejects[0]["reason"] == "ocr_required"
    assert rejects[0]["position"]["page"] == 1
    assert rejects[0]["page_kind"] == "scanned"
    assert rejects[0]["recommended_backend"] == "paddleocr"
    assert rejects[0]["fallback_backends"] == ["qwen_vl"]


class _FakeMinerUClient:
    provider = "mineru"
    base_url = "http://mineru.test:8000"
    backend = "pipeline"

    def __init__(self, *, fail: bool = False) -> None:
        self.fail = fail
        self.calls = 0

    def health(self) -> MinerUHealth:
        return MinerUHealth(status="healthy", version="3.0.9", protocol_version=1)

    def parse_pdf(self, raw_bytes: bytes, *, filename: str) -> MinerUParseResult:
        self.calls += 1
        if self.fail:
            raise MinerUServiceError("MinerU 测试故障")
        return self._result()

    def _result(self) -> MinerUParseResult:
        raw_response = {
            "task_id": "mineru-task-1",
            "status": "completed",
            "backend": "pipeline",
            "version": "3.0.9",
            "results": {"data": {"md_content": "扫描页文本"}},
        }
        return MinerUParseResult(
            task_id="mineru-task-1",
            backend="pipeline",
            version="3.0.9",
            blocks=(
                MinerUPageBlock(
                    page=1,
                    text="扫描页文本",
                    bbox=(10.0, 20.0, 210.0, 60.0),
                    coordinate_space="image_pixels",
                    confidence=0.96,
                    element_type="text",
                ),
            ),
            raw_response=raw_response,
        )

    def parse_response(self, raw_response: dict) -> MinerUParseResult:
        return self._result()


def test_pdf_parser_uses_mineru_for_scanned_page_and_keeps_raw_result(tmp_path: Path) -> None:
    store = ArtifactStore(root=str(tmp_path))
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    data = bytes(pdf.output())
    art = store.write_raw(
        task_id="task-mineru", source_id="file-1", data=data,
        uri="upload://scan.pdf", media_type="application/pdf", ext="pdf",
    )
    client = _FakeMinerUClient()
    parser = PdfParser(
        mineru_client=client,
        artifact_store=store,
        use_remote_ocr=True,
    )

    records, rejects = parser.parse(art, data)

    assert rejects == []
    assert len(records) == 1
    assert records[0].data["text"] == "扫描页文本"
    element = DocumentElement.model_validate(records[0].data["elements"][0])
    assert element.extractor == "mineru"
    assert element.extractor_version == "3.0.9"
    assert element.bbox is not None
    assert element.bbox.coordinate_space == "image_pixels"
    assert element.raw_result_ref
    assert store.resolve_path(element.raw_result_ref).exists()
    assert records[0].meta["route"]["actual_backend"] == "mineru:pipeline"
    assert client.calls == 1


def test_pdf_parser_keeps_ocr_required_when_mineru_fails(tmp_path: Path) -> None:
    store = ArtifactStore(root=str(tmp_path))
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    data = bytes(pdf.output())
    art = store.write_raw(
        task_id="task-mineru-fail", source_id="file-1", data=data,
        uri="upload://scan.pdf", media_type="application/pdf", ext="pdf",
    )
    parser = PdfParser(
        mineru_client=_FakeMinerUClient(fail=True),
        artifact_store=store,
        use_remote_ocr=True,
    )

    records, rejects = parser.parse(art, data)

    assert records == []
    assert len(rejects) == 1
    assert rejects[0]["reason"] == "ocr_required"
    assert rejects[0]["ocr_attempted"] is True
    assert "MinerU 测试故障" in rejects[0]["ocr_error"]


def test_pdf_parser_reuses_immutable_mineru_cache(tmp_path: Path) -> None:
    store = ArtifactStore(root=str(tmp_path))
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    data = bytes(pdf.output())
    art = store.write_raw(
        task_id="task-mineru-cache", source_id="file-1", data=data,
        uri="upload://scan.pdf", media_type="application/pdf", ext="pdf",
    )
    client = _FakeMinerUClient()
    parser = PdfParser(
        mineru_client=client,
        artifact_store=store,
        use_remote_ocr=True,
    )

    first_records, first_rejects = parser.parse(art, data)
    second_records, second_rejects = parser.parse(art, data)

    assert first_rejects == second_rejects == []
    assert first_records[0].data == second_records[0].data
    assert client.calls == 1
    assert first_records[0].meta["route"]["cache_hit"] is False
    assert second_records[0].meta["route"]["cache_hit"] is True


def test_pdf_parser_never_calls_mineru_for_digital_pdf(tmp_path: Path) -> None:
    store = ArtifactStore(root=str(tmp_path))
    data = _make_pdf_bytes(["Digital content"])
    art = store.write_raw(
        task_id="task-digital", source_id="file-1", data=data,
        uri="upload://digital.pdf", media_type="application/pdf", ext="pdf",
    )
    client = _FakeMinerUClient(fail=True)
    parser = PdfParser(
        mineru_client=client,
        artifact_store=store,
        use_remote_ocr=True,
    )

    records, rejects = parser.parse(art, data)

    assert len(records) == 1
    assert rejects == []
    assert client.calls == 0


def test_pdf_parser_falls_back_to_secondary_document_service(tmp_path: Path) -> None:
    store = ArtifactStore(root=str(tmp_path))
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    data = bytes(pdf.output())
    art = store.write_raw(
        task_id="task-document-fallback", source_id="file-1", data=data,
        uri="upload://scan.pdf", media_type="application/pdf", ext="pdf",
    )
    primary = _FakeMinerUClient(fail=True)
    primary.provider = "paddleocr_vl"
    primary.base_url = "http://paddle.test:8080"
    primary.backend = "PaddleOCR-VL-1.6"
    secondary = _FakeMinerUClient()
    parser = PdfParser(
        document_clients=[primary, secondary],
        artifact_store=store,
        use_remote_ocr=True,
    )

    records, rejects = parser.parse(art, data)

    assert rejects == []
    assert len(records) == 1
    assert primary.calls == 1
    assert secondary.calls == 1
    assert records[0].meta["route"]["actual_backend"] == "mineru:pipeline"


def test_pdf_parser_recovers_failed_document_with_single_page_retry(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """整份请求失败后只重试缺失页，并把服务页码映射回原页。"""
    store = ArtifactStore(root=str(tmp_path))
    from fpdf import FPDF
    from src.config.settings import settings

    pdf = FPDF()
    pdf.add_page()
    data = bytes(pdf.output())
    art = store.write_raw(
        task_id="task-page-retry",
        source_id="file-1",
        data=data,
        uri="upload://scan.pdf",
        media_type="application/pdf",
        ext="pdf",
    )

    class FlakyClient(_FakeMinerUClient):
        def parse_pdf(self, raw_bytes: bytes, *, filename: str) -> MinerUParseResult:
            self.calls += 1
            if self.calls == 1:
                raise MinerUServiceError("整份解析瞬时故障")
            return self._result()

    monkeypatch.setattr(settings, "document_parser_page_retries", 1)
    monkeypatch.setattr(settings, "document_parser_retry_backoff_seconds", 0)
    client = FlakyClient()
    parser = PdfParser(
        document_clients=[client],
        artifact_store=store,
        use_remote_ocr=True,
    )

    records, rejects = parser.parse(art, data)

    assert rejects == []
    assert client.calls == 2
    assert records[0].meta["position"]["page"] == 1
    assert records[0].meta["route"]["actual_backend"] == "mineru:pipeline"


def test_pdf_parser_uses_paddle_to_enrich_mineru_table_page(tmp_path: Path) -> None:
    store = ArtifactStore(root=str(tmp_path))
    from fpdf import FPDF

    class MinerUTableClient(_FakeMinerUClient):
        def _result(self) -> MinerUParseResult:
            result = super()._result()
            return MinerUParseResult(
                task_id=result.task_id,
                backend=result.backend,
                version=result.version,
                blocks=(
                    MinerUPageBlock(
                        page=1,
                        text="项目 金额",
                        bbox=(10.0, 20.0, 210.0, 160.0),
                        coordinate_space="normalized_1000",
                        confidence=0.8,
                        element_type="table",
                    ),
                ),
                raw_response=result.raw_response,
            )

    class PaddleTableClient(MinerUTableClient):
        provider = "paddleocr_vl"
        base_url = "http://paddle.test:18081"
        backend = "PaddleOCR-VL-1.6"

        def _result(self) -> MinerUParseResult:
            result = super()._result()
            return MinerUParseResult(
                task_id=result.task_id,
                backend=self.backend,
                version="1.6",
                blocks=(
                    MinerUPageBlock(
                        page=1,
                        text="|项目|金额|\n|A|100|",
                        bbox=(10.0, 20.0, 210.0, 160.0),
                        coordinate_space="normalized_1000",
                        confidence=0.98,
                        element_type="table",
                    ),
                ),
                raw_response=result.raw_response,
                provider=self.provider,
            )

    pdf = FPDF()
    pdf.add_page()
    data = bytes(pdf.output())
    art = store.write_raw(
        task_id="task-table-enrich",
        source_id="file-1",
        data=data,
        uri="upload://table.pdf",
        media_type="application/pdf",
        ext="pdf",
    )
    mineru = MinerUTableClient()
    paddle = PaddleTableClient()
    parser = PdfParser(
        document_clients=[mineru, paddle],
        artifact_store=store,
        use_remote_ocr=True,
    )

    records, rejects = parser.parse(art, data)

    assert rejects == []
    assert mineru.calls == 1
    assert paddle.calls == 1
    table_elements = [
        DocumentElement.model_validate(item)
        for item in records[0].data["elements"]
        if item["element_type"] == "table"
    ]
    assert {item.extractor for item in table_elements} == {
        "mineru",
        "paddleocr_vl",
    }


def test_docx_parser_extracts_paragraphs(tmp_path: Path):
    store = ArtifactStore(root=str(tmp_path))
    data = _make_docx_bytes(["First paragraph", "Second paragraph", ""])
    art = store.write_raw(
        task_id="task-1", source_id="file-1", data=data,
        uri="upload://data.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ext="docx",
    )

    parser = OfficeParser()
    records, rejects = parser.parse(art, data)

    assert len(records) == 2  # 空段落不计
    assert records[0].data["text"] == "First paragraph"
    assert records[0].meta["parser"] == "docx"
    assert records[0].meta["position"]["index"] == 0
    first = DocumentElement.model_validate(records[0].data["elements"][0])
    second = DocumentElement.model_validate(records[1].data["elements"][0])
    assert first.element_type.value == "paragraph"
    assert first.extractor == "python-docx"
    assert first.metadata["location"] == {
        "kind": "docx_paragraph",
        "paragraph": 1,
    }
    assert first.reading_order == 0
    assert second.reading_order == 1
    assert first.bbox is None
    assert OfficeParser().parse(art, data)[0][0].data["elements"][0] == (
        records[0].data["elements"][0]
    )
    assert rejects == []


def test_docx_parser_extracts_table_rows(tmp_path: Path):
    from docx import Document

    store = ArtifactStore(root=str(tmp_path))
    doc = Document()
    doc.add_paragraph("Intro")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "id"
    table.rows[0].cells[1].text = "name"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "Alice"
    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    art = store.write_raw(
        task_id="task-1", source_id="file-1", data=data,
        uri="upload://data.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ext="docx",
    )

    parser = OfficeParser()
    records, rejects = parser.parse(art, data)

    # 1 段落 + 1 数据行（表头不计为数据记录）
    assert len(records) == 2
    table_records = [r for r in records if r.meta["parser"] == "docx_table"]
    assert len(table_records) == 1
    assert table_records[0].data["id"] == "1"
    assert table_records[0].data["name"] == "Alice"
    element = DocumentElement.model_validate(
        table_records[0].data["elements"][0]
    )
    assert element.element_type.value == "table"
    assert element.text == "id：1；name：Alice"
    assert element.metadata["location"] == {
        "kind": "docx_table_row",
        "table": 1,
        "row": 1,
    }


def test_docx_parser_preserves_paragraph_table_paragraph_order(tmp_path: Path):
    from docx import Document

    store = ArtifactStore(root=str(tmp_path))
    doc = Document()
    doc.add_paragraph("付款条款", style="Heading 1")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "金额"
    table.rows[1].cells[0].text = "首付款"
    table.rows[1].cells[1].text = "60%"
    doc.add_paragraph("交付条款")
    buffer = io.BytesIO()
    doc.save(buffer)
    data = buffer.getvalue()
    artifact = store.write_raw(
        task_id="task-order",
        source_id="file-1",
        data=data,
        uri="upload://contract.docx",
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        ext="docx",
    )

    records, rejects = OfficeParser().parse(artifact, data)
    elements = [
        DocumentElement.model_validate(record.data["elements"][0])
        for record in records
    ]

    assert rejects == []
    assert [item.reading_order for item in elements] == [0, 1, 2]
    assert [item.text for item in elements] == [
        "付款条款",
        "项目：首付款；金额：60%",
        "交付条款",
    ]
    assert elements[0].element_type.value == "heading"
