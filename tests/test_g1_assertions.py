# -*- coding: utf-8 -*-
from __future__ import annotations

import importlib.util
import json
from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook


ASSERTIONS_PATH = (
    Path(__file__).resolve().parents[1]
    / "evals/generalization-g1/assertions.py"
)
SPEC = importlib.util.spec_from_file_location("g1_assertions", ASSERTIONS_PATH)
assert SPEC is not None and SPEC.loader is not None
ASSERTIONS = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(ASSERTIONS)


def test_formal_mode_rejects_historical_weak_rules() -> None:
    fixtures = {
        "cases": [
            {"id": "weak", "assert_rule": "pdf_page3_table"},
            {"id": "strong", "assert_rule": "xlsx_scene_planning"},
        ]
    }

    assert ASSERTIONS.formal_assertion_gaps(fixtures) == (
        "weak 使用非正式强度断言：pdf_page3_table",
    )


def test_d2_rejects_long_json_unrelated_to_source(tmp_path: Path) -> None:
    source = tmp_path / "requirements.docx"
    document = Document()
    document.add_heading("功能需求模块", level=1)
    for module in ("订单创建", "订单审批", "订单跟踪"):
        document.add_paragraph(module)
    document.save(source)
    candidate = tmp_path / "candidate.json"
    candidate.write_text(
        json.dumps(
            {"modules": ["天气预报", "音乐播放", "旅游推荐"] * 10},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    case = {"id": "D2-negative", "assert_rule": "docx_po_modules"}

    with pytest.raises(AssertionError, match="源中无出处"):
        ASSERTIONS.run_assert(
            case,
            candidate,
            source_resolver=lambda _case: [(source, {"sha256": "test"})],
        )
    candidate.write_text(
        json.dumps(
            {"modules": ["订单创建", "订单审批", "订单跟踪"]},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    ASSERTIONS.run_assert(
        case,
        candidate,
        source_resolver=lambda _case: [(source, {"sha256": "test"})],
    )


def test_d4_rejects_three_rows_with_wrong_business_values(tmp_path: Path) -> None:
    source = tmp_path / "agents.docx"
    document = Document()
    table = document.add_table(rows=1, cols=3)
    for index, value in enumerate(("智能体", "应用场景数", "核心价值")):
        table.rows[0].cells[index].text = value
    for values in (
        ("需求分析智能体", "5", "缩短需求澄清周期"),
        ("测试设计智能体", "3", "提高测试覆盖率"),
        ("研发助手智能体", "8", "减少重复编码"),
    ):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    document.save(source)
    candidate = tmp_path / "candidate.csv"
    candidate.write_text(
        "智能体,应用场景数,核心价值\n天气助手,99,查询天气\n音乐助手,88,播放音乐\n旅游助手,77,推荐景点\n",
        encoding="utf-8",
    )
    case = {"id": "D4-negative", "assert_rule": "docx_agents_scenarios"}

    with pytest.raises(AssertionError, match="与源不一致"):
        ASSERTIONS.run_assert(
            case,
            candidate,
            source_resolver=lambda _case: [(source, {"sha256": "test"})],
        )
    candidate.write_text(
        "智能体,应用场景数,核心价值\n"
        "测试设计智能体,3,提高测试覆盖率\n"
        "需求分析智能体,5,缩短需求澄清周期\n"
        "研发助手智能体,8,减少重复编码\n",
        encoding="utf-8",
    )
    ASSERTIONS.run_assert(
        case,
        candidate,
        source_resolver=lambda _case: [(source, {"sha256": "test"})],
    )


def test_x2_rejects_many_rows_with_wrong_scenes(tmp_path: Path) -> None:
    source = tmp_path / "planning.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "大师场景规划"
    sheet.append(["大师（一级）", "智能体（二级）", "应用场景（三级）"])
    for index in range(1, 21):
        sheet.append(["研发大师", f"智能体{index}", f"真实场景{index}"])
    workbook.save(source)
    candidate = tmp_path / "candidate.csv"
    candidate.write_text(
        "大师,智能体,应用场景\n"
        + "\n".join(
            f"无关大师,无关智能体{index},虚构场景{index}" for index in range(1, 21)
        ),
        encoding="utf-8",
    )
    case = {"id": "X2-negative", "assert_rule": "xlsx_scene_planning"}

    with pytest.raises(AssertionError, match="与源不一致"):
        ASSERTIONS.run_assert(
            case,
            candidate,
            source_resolver=lambda _case: [(source, {"sha256": "test"})],
        )
    candidate.write_text(
        "应用场景\n" + "\n".join(f"真实场景{index}" for index in range(20, 0, -1)),
        encoding="utf-8",
    )
    ASSERTIONS.run_assert(
        case,
        candidate,
        source_resolver=lambda _case: [(source, {"sha256": "test"})],
    )


def test_x3_rejects_many_rows_that_are_not_2026_new_scenes(tmp_path: Path) -> None:
    source = tmp_path / "delivery-plan.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "V1.0 26年场景交付计划表"
    sheet.append(["智能体（二级）", "应用场景（三级）", "2026", "26年建设类型"])
    for index in range(1, 11):
        sheet.append([f"智能体{index}", f"新增场景{index}", "A2", "新增"])
    sheet.append(["已有智能体", "迭代场景", "A3", "迭代"])
    workbook.save(source)
    candidate = tmp_path / "candidate.csv"
    candidate.write_text(
        "应用场景,建设类型\n"
        + "\n".join(f"虚构场景{index},新增" for index in range(1, 11)),
        encoding="utf-8",
    )
    case = {"id": "X3-negative", "assert_rule": "xlsx_plan2026_new"}

    with pytest.raises(AssertionError, match="与源不一致"):
        ASSERTIONS.run_assert(
            case,
            candidate,
            source_resolver=lambda _case: [(source, {"sha256": "test"})],
        )
    candidate.write_text(
        "应用场景\n" + "\n".join(f"新增场景{index}" for index in range(10, 0, -1)),
        encoding="utf-8",
    )
    ASSERTIONS.run_assert(
        case,
        candidate,
        source_resolver=lambda _case: [(source, {"sha256": "test"})],
    )


def test_f1_rejects_long_text_without_source_money_or_time_facts(tmp_path: Path) -> None:
    source = tmp_path / "requirements.docx"
    document = Document()
    document.add_paragraph("项目预算为100万元，须在2026年6月30日前上线。")
    document.add_paragraph("上线后稳定运行3个月再进行验收。")
    document.save(source)
    candidate = tmp_path / "candidate.txt"
    candidate.write_text("这是一段完全无关的长文本。" * 30, encoding="utf-8")
    case = {"id": "F1-negative", "assert_rule": "docx_money_time_extract"}

    with pytest.raises(AssertionError, match="缺少源中的钱或时间事实"):
        ASSERTIONS.run_assert(
            case,
            candidate,
            source_resolver=lambda _case: [(source, {"sha256": "test"})],
        )
    candidate.write_text(
        "预算：100万元；上线日期：2026年6月30日；验收前稳定运行：3个月。",
        encoding="utf-8",
    )
    ASSERTIONS.run_assert(
        case,
        candidate,
        source_resolver=lambda _case: [(source, {"sha256": "test"})],
    )


def test_f2_rejects_nonempty_but_wrong_2026_plans(tmp_path: Path) -> None:
    source = tmp_path / "plans.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "大师建设计划进展(副本)"
    sheet.append(["AI大师（一级）", "2026年", "2027年"])
    sheet.append(["数智软件大师", "A2", "A3"])
    sheet.append(["产品策划大师", "A1", "A2"])
    workbook.save(source)
    candidate = tmp_path / "candidate.csv"
    candidate.write_text(
        "AI大师,2026年\n天气大师,A5\n音乐大师,A4\n",
        encoding="utf-8",
    )
    case = {"id": "F2-negative", "assert_rule": "xlsx_2026_plan"}

    with pytest.raises(AssertionError, match="与源不一致"):
        ASSERTIONS.run_assert(
            case,
            candidate,
            source_resolver=lambda _case: [(source, {"sha256": "test"})],
        )
    candidate.write_text(
        "AI大师,2026年\n产品策划大师,A1\n数智软件大师,A2\n",
        encoding="utf-8",
    )
    ASSERTIONS.run_assert(
        case,
        candidate,
        source_resolver=lambda _case: [(source, {"sha256": "test"})],
    )
