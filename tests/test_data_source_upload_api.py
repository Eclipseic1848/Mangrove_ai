# -*- coding: utf-8 -*-
"""数据源上传 API 测试（Phase 2 Task 3）。"""
from __future__ import annotations

import io
from pathlib import Path

import pytest
from fastapi import FastAPI
from fastapi.testclient import TestClient

from src.api.auth import get_current_user
from src.api.routes.data_sources import router
from src.config.settings import settings


def _make_client(tmp_path: Path, monkeypatch, *, max_bytes: int = 1024, user_id: str = "user-a") -> TestClient:
    monkeypatch.setattr(settings, "data_prep_upload_root", str(tmp_path))
    monkeypatch.setattr(settings, "data_prep_max_upload_bytes", max_bytes)
    app = FastAPI()
    app.include_router(router)
    app.dependency_overrides[get_current_user] = lambda: {"user_id": user_id}
    return TestClient(app)


def test_upload_returns_item_without_sensitive_fields(tmp_path: Path, monkeypatch):
    client = _make_client(tmp_path, monkeypatch)
    data = b"id,name\n1,Alice\n2,Bob\n"
    response = client.post(
        "/api/data-sources/uploads",
        files={"file": ("data.csv", data, "text/csv")},
    )

    assert response.status_code == 200, response.text
    body = response.json()
    assert body["original_name"] == "data.csv"
    assert body["size_bytes"] == len(data)
    assert "storage_path" not in body
    assert "user_id" not in body
    assert body["upload_id"]


def test_other_user_cannot_access_upload(tmp_path: Path, monkeypatch):
    client_a = _make_client(tmp_path, monkeypatch, user_id="user-a")
    resp = client_a.post(
        "/api/data-sources/uploads",
        files={"file": ("data.csv", b"id\n1\n", "text/csv")},
    )
    upload_id = resp.json()["upload_id"]

    client_b = _make_client(tmp_path, monkeypatch, user_id="user-b")
    assert client_b.get(f"/api/data-sources/uploads/{upload_id}").status_code == 404
    assert client_b.delete(f"/api/data-sources/uploads/{upload_id}").status_code == 404


def test_owner_can_get_and_delete(tmp_path: Path, monkeypatch):
    client = _make_client(tmp_path, monkeypatch, user_id="user-a")
    resp = client.post(
        "/api/data-sources/uploads",
        files={"file": ("data.csv", b"id\n1\n", "text/csv")},
    )
    upload_id = resp.json()["upload_id"]

    assert client.get(f"/api/data-sources/uploads/{upload_id}").status_code == 200
    assert client.delete(f"/api/data-sources/uploads/{upload_id}").status_code == 200
    assert client.get(f"/api/data-sources/uploads/{upload_id}").status_code == 404


def test_oversized_upload_returns_413(tmp_path: Path, monkeypatch):
    client = _make_client(tmp_path, monkeypatch, max_bytes=64)
    response = client.post(
        "/api/data-sources/uploads",
        files={"file": ("big.csv", b"x" * 1024, "text/csv")},
    )

    assert response.status_code == 413


def test_mismatched_magic_returns_413(tmp_path: Path, monkeypatch):
    """PNG 魔数 + .csv 扩展名被拒绝。"""
    client = _make_client(tmp_path, monkeypatch, max_bytes=1024)
    png_header = b"\x89PNG\r\n\x1a\n" + b"\x00" * 32
    response = client.post(
        "/api/data-sources/uploads",
        files={"file": ("fake.csv", png_header, "text/csv")},
    )

    assert response.status_code == 413


def test_docx_upload_has_structured_preview(tmp_path: Path, monkeypatch):
    """上传后的 DOCX 可立即预览，元素 ID 可重复解析且保持稳定。"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("付款条件：验收后 30 日内付款")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "订单号"
    table.rows[0].cells[1].text = "金额"
    table.rows[1].cells[0].text = "PO-001"
    table.rows[1].cells[1].text = "1000 元"
    buffer = io.BytesIO()
    doc.save(buffer)

    client = _make_client(tmp_path, monkeypatch, max_bytes=1024 * 1024)
    upload = client.post(
        "/api/data-sources/uploads",
        files={
            "file": (
                "contract.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload.status_code == 200, upload.text
    upload_id = upload.json()["upload_id"]

    first = client.get(
        f"/api/data-sources/uploads/{upload_id}/document-preview"
    )
    second = client.get(
        f"/api/data-sources/uploads/{upload_id}/document-preview"
    )

    assert first.status_code == 200, first.text
    assert first.json()["status"] == "ready"
    assert [item["element_type"] for item in first.json()["elements"]] == [
        "paragraph",
        "table",
    ]
    assert first.json()["elements"][0]["metadata"]["location"]["paragraph"] == 1
    assert first.json()["elements"][1]["metadata"]["location"]["table"] == 1
    assert [item["element_id"] for item in first.json()["elements"]] == [
        item["element_id"] for item in second.json()["elements"]
    ]


def test_document_preview_rejects_non_docx(tmp_path: Path, monkeypatch):
    client = _make_client(tmp_path, monkeypatch)
    upload = client.post(
        "/api/data-sources/uploads",
        files={"file": ("data.csv", b"id\n1\n", "text/csv")},
    )
    upload_id = upload.json()["upload_id"]

    response = client.get(
        f"/api/data-sources/uploads/{upload_id}/document-preview"
    )

    assert response.status_code == 415
