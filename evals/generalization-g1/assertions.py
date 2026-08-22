# -*- coding: utf-8 -*-
"""G1 泛化集确定性断言规则。

纪律：
- 能可靠推导的期望值从源文件动态计算；部分规则暂用固定结构/长度/行数阈值，只具诊断强度；
- 扫描件 PDF 类先做结构断言，真实试运行人工核验后再回填锚点（回填后重新冻结）；
- 每条规则签名：assert_<rule>(case, candidate_path) -> None，失败抛 AssertionError。
"""
from __future__ import annotations

import csv
import hashlib
import json
from pathlib import Path
import re
import shutil
import tempfile
import unicodedata
from collections import Counter
from collections.abc import Callable

from docx import Document
from openpyxl import load_workbook
import pdfplumber

EVALS_ROOT = Path(__file__).resolve().parent
PROJECT_ROOT = EVALS_ROOT.parents[1]
UPLOAD_ROOT = PROJECT_ROOT / "data" / "uploads"
SourceResolver = Callable[[dict], list[tuple[Path, dict]]]


def resolve_sources(case: dict) -> list[tuple[Path, dict]]:
    """按内容 sha256 精确匹配上传对象（同名多对象用哈希消歧）。

    上传对象文件名是 upload_id（32 hex），内容哈希在 .meta 里；必须读取
    meta 后按内容哈希匹配，并复验对象文件哈希，双校验后才可用于评测。
    """
    found: list[tuple[Path, dict]] = []
    for ref in case["sources"]:
        candidates = []
        for meta_path in UPLOAD_ROOT.glob("*/objects/*.meta"):
            metadata = json.loads(meta_path.read_text(encoding="utf-8"))
            if metadata.get("sha256") == ref["sha256"]:
                candidates.append((meta_path, metadata))
        if not candidates:
            raise AssertionError(
                f"{case['id']}: 语料 {ref['original_name']} 按内容哈希未命中"
            )
        meta_path, metadata = candidates[0]
        if metadata.get("original_name") != ref["original_name"]:
            raise AssertionError(
                f"{case['id']}: 哈希命中但文件名不一致："
                f"{metadata.get('original_name')} != {ref['original_name']}"
            )
        source = meta_path.with_suffix("")
        actual = hashlib.sha256(source.read_bytes()).hexdigest()
        if actual != ref["sha256"]:
            raise AssertionError(f"{case['id']}: 语料对象哈希与冻结清单不一致")
        found.append((source, metadata))
    return found


def normalized(value: object) -> str:
    return re.sub(r"\s+", "", unicodedata.normalize("NFKC", str(value or "")))


def read_candidate_csv(path: Path) -> list[list[str]]:
    with path.open(encoding="utf-8-sig", errors="replace", newline="") as handle:
        rows = list(csv.reader(handle))
    if not rows or all(not any(cell.strip() for cell in row) for row in rows):
        raise AssertionError("候选 CSV 为空")
    return rows


def candidate_text(path: Path) -> str:
    content = path.read_text(encoding="utf-8-sig", errors="replace")
    if not content.strip():
        raise AssertionError("候选文本为空")
    return content


def candidate_json(path: Path) -> object:
    return json.loads(path.read_text(encoding="utf-8-sig", errors="replace"))


def _json_leaf_strings(value: object) -> list[str]:
    if isinstance(value, dict):
        return [item for child in value.values() for item in _json_leaf_strings(child)]
    if isinstance(value, list):
        return [item for child in value for item in _json_leaf_strings(child)]
    if isinstance(value, str) and len(normalized(value)) >= 2:
        return [value]
    return []


def _pdf_text(source: Path) -> str:
    with pdfplumber.open(source) as pdf:
        return "\n".join(
            page.extract_text() or "" for page in pdf.pages
        )


def _docx_text(source: Path) -> str:
    document = Document(source)
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text.strip() for cell in row.cells if cell.text.strip())
    return "\n".join(parts)


def _docx_tables(source: Path) -> list[list[list[str]]]:
    document = Document(source)
    return [
        [[cell.text.strip() for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]


def _xlsx_sheets(source: Path) -> dict[str, list[list[str]]]:
    """openpyxl 按扩展名判定格式，先复制成带 .xlsx 名的临时文件。"""
    tmp = Path(tempfile.mkdtemp(prefix="g1-xlsx-")) / "book.xlsx"
    shutil.copyfile(source, tmp)
    workbook = load_workbook(tmp, read_only=True, data_only=True)
    try:
        return {
            name: [
                ["" if cell is None else str(cell) for cell in row]
                for row in workbook[name].iter_rows(values_only=True)
            ]
            for name in workbook.sheetnames
        }
    finally:
        workbook.close()
        shutil.rmtree(tmp.parent, ignore_errors=True)


def _find_header(rows: list[list[str]], *needles: str) -> int:
    for index, row in enumerate(rows):
        flat = [normalized(cell) for cell in row]
        if all(any(needle in cell for cell in flat) for needle in needles):
            return index
    raise AssertionError(f"候选缺少表头：{'/'.join(needles)}")


def _nonempty_rows(rows: list[list[str]]) -> list[list[str]]:
    return [row for row in rows if any(cell.strip() for cell in row)]


# ---------------------------------------------------------------- 规则
def pdf_workload_detail(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    header = _find_header(rows, "姓名") if any("姓名" in normalized(c) for r in rows for c in r) else None
    data = _nonempty_rows(rows[header + 1 :] if header is not None else rows)
    if len(data) < 3:
        raise AssertionError(f"明细行数异常（期望 ≥3，实际 {len(data)}）")


def pdf_workload_summary(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    data = _nonempty_rows(rows[1:] if rows else [])
    if len(data) < 2:
        raise AssertionError(f"汇总行数异常（期望 ≥2，实际 {len(data)}）")


def pdf_workload_detail_second_source(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    if len(_nonempty_rows(rows)) < 3:
        raise AssertionError("第二张核算表明细行数异常")


def pdf_workload_merge_three(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    data = _nonempty_rows(rows)
    if len(data) < 9:
        raise AssertionError(f"三张表合并后行数异常（期望 ≥9，实际 {len(data)}）")


def pdf_page3_table(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    data = _nonempty_rows(rows)
    if not data:
        raise AssertionError("第 3 页表格候选为空")


def pdf_first_document(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    data = _nonempty_rows(rows)
    if not data:
        raise AssertionError("第一张单据候选为空")


def pdf_all_documents(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    data = _nonempty_rows(rows)
    if len(data) < 5:
        raise AssertionError(f"全量单据行数异常（期望 ≥5，实际 {len(data)}）")


def pdf_function_modules(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    data = _nonempty_rows(rows)
    source_text = _pdf_text(resolve_sources(case)[0][0])
    if len(data) < 3:
        raise AssertionError(f"功能模块清单行数异常（期望 ≥3，实际 {len(data)}）")
    # 每个候选项必须能在源中找到出处（防止凭空编造）
    for row in data:
        if not any(normalized(cell) and normalized(cell) in normalized(source_text) for cell in row):
            raise AssertionError(f"候选行在源中无出处：{row}")


def pdf_plan_topics(case: dict, candidate: Path) -> None:
    content = candidate_text(candidate)
    source_text = normalized(_pdf_text(resolve_sources(case)[0][0]))
    for keyword in ("数据集", "智能体", "AICoding", "座舱"):
        if keyword not in source_text:
            continue
        if keyword not in content:
            raise AssertionError(f"候选缺少源中主题：{keyword}")


def pdf_miniapp_quiz_steps(case: dict, candidate: Path) -> None:
    content = candidate_text(candidate)
    if "题库" not in content:
        raise AssertionError("候选未覆盖题库模块")
    if len(content) < 100:
        raise AssertionError("候选过短")


def docx_core_indicators(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    tables = _docx_tables(resolve_sources(case)[0][0])
    # 目标表 = 指标类表（表头含「指标名称」或「定义」）
    target = next(
        (t for t in tables if any("指标名称" in normalized(cell) for row in t for cell in row)),
        None,
    )
    if target is None:
        raise AssertionError("源中未找到指标表")
    expected_rows = sum(
        1 for row in target[1:] if any(cell.strip() for cell in row)
    )
    header = _find_header(rows, "指标")
    data = _nonempty_rows(rows[header + 1 :])
    if len(data) != expected_rows:
        raise AssertionError(
            f"指标行数与源不一致（期望 {expected_rows}，实际 {len(data)}）"
        )


def docx_po_modules(
    case: dict,
    candidate: Path,
    *,
    source_resolver: SourceResolver = resolve_sources,
) -> None:
    obj = candidate_json(candidate)
    if not isinstance(obj, list) and not isinstance(obj, dict):
        raise AssertionError("候选 JSON 结构异常")
    source_text = normalized(_docx_text(source_resolver(case)[0][0]))
    values = _json_leaf_strings(obj)
    if len(values) < 3:
        raise AssertionError("候选 JSON 缺少至少三个功能模块值")
    unsupported = [value for value in values if normalized(value) not in source_text]
    if unsupported:
        raise AssertionError(f"候选模块在源中无出处：{unsupported[:3]}")


def docx_region_elements(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    data = _nonempty_rows(rows)
    if len(data) < 5:
        raise AssertionError(f"区域元素行数异常（期望 ≥5，实际 {len(data)}）")


def docx_agents_scenarios(
    case: dict,
    candidate: Path,
    *,
    source_resolver: SourceResolver = resolve_sources,
) -> None:
    rows = read_candidate_csv(candidate)
    tables = _docx_tables(source_resolver(case)[0][0])
    target = next(
        (
            table
            for table in tables
            if table
            and all(
                any(needle in normalized(cell) for cell in table[0])
                for needle in ("智能体", "应用场景", "核心价值")
            )
        ),
        None,
    )
    if target is None:
        raise AssertionError("源中未找到智能体应用场景表")

    def select_columns(table_rows: list[list[str]], header_index: int) -> list[tuple[str, str, str]]:
        header = [normalized(cell) for cell in table_rows[header_index]]
        indexes = [
            next(i for i, cell in enumerate(header) if needle in cell)
            for needle in ("智能体", "应用场景", "核心价值")
        ]
        selected = []
        for row in table_rows[header_index + 1 :]:
            values = tuple(
                normalized(row[index]) if index < len(row) else ""
                for index in indexes
            )
            if any(values):
                selected.append(values)
        return selected

    expected = select_columns(target, 0)
    candidate_header = _find_header(rows, "智能体", "应用场景", "核心价值")
    actual = select_columns(rows, candidate_header)
    if Counter(actual) != Counter(expected):
        raise AssertionError(
            f"智能体应用场与源不一致（期望 {len(expected)} 行，实际 {len(actual)} 行）"
        )


def docx_quality_indicators(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    tables = _docx_tables(resolve_sources(case)[0][0])
    target = next(
        (t for t in tables if any("核心指标" in normalized(cell) for row in t for cell in row)),
        None,
    )
    if target is None:
        raise AssertionError("源中未找到核心指标表")
    expected_rows = sum(
        1 for row in target[1:] if any(cell.strip() for cell in row)
    )
    header = _find_header(rows, "指标")
    data = _nonempty_rows(rows[header + 1 :])
    if len(data) != expected_rows:
        raise AssertionError(
            f"核心指标行数与源不一致（期望 {expected_rows}，实际 {len(data)}）"
        )


def docx_otd_schedule(case: dict, candidate: Path) -> None:
    content = candidate_text(candidate)
    if len(content) < 80:
        raise AssertionError("候选过短")


def xlsx_sheet2_filter(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    data = _nonempty_rows(rows)
    sheets = _xlsx_sheets(resolve_sources(case)[0][0])
    expected_sheet = sheets["工作表2"]
    header_index = _find_header(expected_sheet, "计划交付月份")
    # 期望 = 1 月行，且只留应用场景（三级）与目标成熟度两列；
    # 表头列按关键词定位（NFKC 会把全角括号归一化为半角，不能按精确名找）
    header = [normalized(cell) for cell in expected_sheet[header_index]]
    month_idx = next(
        i for i, cell in enumerate(header) if "计划交付月份" in cell
    )
    scene_idx = next(
        i for i, cell in enumerate(header) if "应用场景" in cell and "三级" in cell
    )
    maturity_idx = next(
        i for i, cell in enumerate(header) if "成熟度" in cell
    )
    expected = [
        (normalized(row[scene_idx]), normalized(row[maturity_idx]))
        for row in expected_sheet[header_index + 1 :]
        if len(row) > month_idx and normalized(row[month_idx]) == "1月"
        and normalized(row[scene_idx])
    ]
    if not expected:
        raise AssertionError("源工作表2 无 1 月记录")
    # 候选表头允许列名变体（如「应用场景,成熟度」而非精确的
    # 「应用场景（三级）,目标成熟度」），按关键词定位列。
    candidate_header = _find_header(rows, "应用场景", "成熟度")
    candidate_cols = [normalized(cell) for cell in rows[candidate_header]]
    scene_idx = next(
        i for i, cell in enumerate(candidate_cols) if "应用场景" in cell
    )
    maturity_idx = next(
        i for i, cell in enumerate(candidate_cols) if "成熟度" in cell
    )
    candidate_data = [
        (normalized(r[scene_idx]), normalized(r[maturity_idx]))
        if max(scene_idx, maturity_idx) < len(r) else ("", "")
        for r in rows[candidate_header + 1 :]
    ]
    candidate_data = [r for r in candidate_data if r[0]]
    if candidate_data != expected:
        raise AssertionError(
            f"筛选结果不一致（期望 {len(expected)} 行，实际 {len(candidate_data)} 行）"
        )


def xlsx_scene_planning(
    case: dict,
    candidate: Path,
    *,
    source_resolver: SourceResolver = resolve_sources,
) -> None:
    rows = read_candidate_csv(candidate)
    sheets = _xlsx_sheets(source_resolver(case)[0][0])
    source_rows = sheets.get("大师场景规划")
    if not source_rows:
        raise AssertionError("源中未找到大师场景规划表")
    source_header = _find_header(source_rows, "应用场景")
    source_columns = [normalized(cell) for cell in source_rows[source_header]]
    source_scene_index = next(
        i for i, cell in enumerate(source_columns) if "应用场景" in cell
    )
    expected = [
        normalized(row[source_scene_index])
        for row in source_rows[source_header + 1 :]
        if source_scene_index < len(row) and normalized(row[source_scene_index])
    ]
    candidate_header = _find_header(rows, "应用场景")
    candidate_columns = [normalized(cell) for cell in rows[candidate_header]]
    candidate_scene_index = next(
        i for i, cell in enumerate(candidate_columns) if "应用场景" in cell
    )
    actual = [
        normalized(row[candidate_scene_index])
        for row in rows[candidate_header + 1 :]
        if candidate_scene_index < len(row) and normalized(row[candidate_scene_index])
    ]
    if Counter(actual) != Counter(expected):
        raise AssertionError(
            f"大师场景规划与源不一致（期望 {len(expected)} 项，实际 {len(actual)} 项）"
        )


def xlsx_plan2026_new(
    case: dict,
    candidate: Path,
    *,
    source_resolver: SourceResolver = resolve_sources,
) -> None:
    rows = read_candidate_csv(candidate)
    sheets = _xlsx_sheets(source_resolver(case)[0][0])
    source_rows = sheets.get("V1.0 26年场景交付计划表")
    if not source_rows:
        raise AssertionError("源中未找到 V1.0 26年场景交付计划表")
    source_header = _find_header(source_rows, "应用场景", "建设类型")
    source_columns = [normalized(cell) for cell in source_rows[source_header]]
    source_scene_index = next(
        i for i, cell in enumerate(source_columns) if "应用场景" in cell
    )
    source_type_index = next(
        i for i, cell in enumerate(source_columns) if "建设类型" in cell
    )
    expected = [
        normalized(row[source_scene_index])
        for row in source_rows[source_header + 1 :]
        if max(source_scene_index, source_type_index) < len(row)
        and normalized(row[source_type_index]) == "新增"
        and normalized(row[source_scene_index])
    ]
    candidate_header = _find_header(rows, "应用场景")
    candidate_columns = [normalized(cell) for cell in rows[candidate_header]]
    candidate_scene_index = next(
        i for i, cell in enumerate(candidate_columns) if "应用场景" in cell
    )
    actual = [
        normalized(row[candidate_scene_index])
        for row in rows[candidate_header + 1 :]
        if candidate_scene_index < len(row) and normalized(row[candidate_scene_index])
    ]
    if Counter(actual) != Counter(expected):
        raise AssertionError(
            f"2026 新增场景与源不一致（期望 {len(expected)} 项，实际 {len(actual)} 项）"
        )


def xlsx_ktm_owner(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    data = _nonempty_rows(rows[1:])  # 跳过候选表头行
    sheets = _xlsx_sheets(resolve_sources(case)[0][0])
    ktm = sheets["KTM"]
    # 源推导：只统计「牵头单位」列含「效率部」的行（其他列出现
    # 「效率部」不算——曾把配合单位列误计入，导致期望虚高）
    owner_idx = next(
        i for i, cell in enumerate(ktm[0]) if "牵头单位" in normalized(cell)
    )
    expected = [
        r for r in ktm[1:]
        if any(cell.strip() for cell in r)
        and owner_idx < len(r)
        and "效率部" in str(r[owner_idx])
    ]
    if not expected:
        raise AssertionError("源 KTM 表中未找到效率部牵头行")
    if len(data) != len(expected):
        raise AssertionError(
            f"效率部工作事项行数与源不一致（期望 {len(expected)}，实际 {len(data)}）"
        )


def csv_sum_by_dept(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    expected = {}
    for row in _read_source_csv(resolve_sources(case)[0][0])[1:]:
        if len(row) >= 2 and normalized(row[0]) and normalized(row[1]):
            expected[normalized(row[0])] = expected.get(normalized(row[0]), 0) + float(row[1])
    header = _find_header(rows, "部门", "金额")
    actual = {}
    for row in rows[header + 1 :]:
        if len(row) >= 2 and normalized(row[0]):
            actual[normalized(row[0])] = float(row[1])
    if actual != expected:
        raise AssertionError(f"部门汇总不一致（期望 {expected}，实际 {actual}）")


def _read_source_csv(source: Path) -> list[list[str]]:
    with source.open(encoding="utf-8-sig", errors="replace", newline="") as handle:
        return list(csv.reader(handle))


def csv_total_amount(case: dict, candidate: Path) -> None:
    content = candidate_text(candidate)
    expected = sum(float(r[1]) for r in _read_source_csv(resolve_sources(case)[0][0])[1:] if len(r) >= 2)
    for value in (str(int(expected)), str(float(expected))):
        if value in content:
            break
    else:
        raise AssertionError(f"候选未包含期望总金额 {expected}")


def csv_filter_value(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    source = _read_source_csv(resolve_sources(case)[0][0])
    expected = [normalized(r[1]) for r in source[1:] if len(r) >= 3 and float(r[2]) > 15]
    actual = [normalized(r[1]) for r in _nonempty_rows(rows[1:] if rows else []) if len(r) >= 2]
    if actual != expected:
        raise AssertionError(f"筛选结果不一致（期望 {expected}，实际 {actual}）")


def csv_garbled_header_sum(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    header = _find_header(rows, "部门", "金额")
    data = _nonempty_rows(rows[header + 1 :])
    # 按列求和：正确输出是恰好一行合计（第一列乱码保留原样或为空均可）
    if len(data) != 1:
        raise AssertionError(f"按列求和应恰好一行合计，实际 {len(data)} 行")
    source = _read_source_csv(resolve_sources(case)[0][0])
    expected_total = sum(float(r[1]) for r in source[1:] if len(r) >= 2)
    actual_total = sum(float(r[1]) for r in data if len(r) >= 2)
    if abs(actual_total - expected_total) > 1e-6:
        raise AssertionError(f"求和与源不一致（期望 {expected_total}，实际 {actual_total}）")


def csv_merge_by_name(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    data = _nonempty_rows(rows[1:])  # 跳过候选表头行
    sources = [_read_source_csv(p) for p, _ in resolve_sources(case)]
    expected_names = {normalized(r[0]) for src in sources for r in src[1:] if r and normalized(r[0])}
    actual_names = {normalized(r[0]) for r in data if r and normalized(r[0])}
    if actual_names != expected_names:
        raise AssertionError(f"合并 name 集合不一致（期望 {expected_names}，实际 {actual_names}）")


def csv_combined_total(case: dict, candidate: Path) -> None:
    content = candidate_text(candidate)
    sources = [_read_source_csv(p) for p, _ in resolve_sources(case)]
    expected = sum(float(r[1]) for src in sources for r in src[1:] if len(r) >= 2)
    if str(int(expected)) not in content and str(float(expected)) not in content:
        raise AssertionError(f"候选未包含总账金额 {expected}")


def docx_two_docs_modules(case: dict, candidate: Path) -> None:
    content = candidate_text(candidate)
    sources = [p for p, _ in resolve_sources(case)]
    if "数智" not in content and "研发" not in content:
        raise AssertionError("候选未覆盖第一份文档")
    if "AI" not in content and "Process" not in content and "官网" not in content:
        raise AssertionError("候选未覆盖第二份文档")


def csv_three_merge(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    data = _nonempty_rows(rows[1:])  # 跳过候选表头行
    sources = [_read_source_csv(p) for p, _ in resolve_sources(case)]
    expected_rows = sum(len(src) - 1 for src in sources)  # 全部数据行，不删不并
    if len(data) != expected_rows:
        raise AssertionError(f"三表合并应保留全部数据行（期望 {expected_rows}，实际 {len(data)}）")
    # 结构检查：部门金额组行与非同构组行都存在（列分组合并）
    row_texts = [normalized(" ".join(row)) for row in data]
    if not any("研发" in text for text in row_texts):
        raise AssertionError("合并结果缺少部门金额组数据")
    if not any("Alice" in text or "Carol" in text for text in row_texts):
        raise AssertionError("合并结果缺少 phase2 组数据")


def docx_money_time_extract(
    case: dict,
    candidate: Path,
    *,
    source_resolver: SourceResolver = resolve_sources,
) -> None:
    content = normalized(candidate_text(candidate))
    source_text = normalized(_docx_text(source_resolver(case)[0][0]))
    fact_pattern = re.compile(
        r"\d+(?:\.\d+)?(?:万|亿)?元"
        r"|20\d{2}年\d{1,2}月\d{1,2}日"
        r"|\d+(?:\.\d+)?(?:个)?(?:工作)?(?:小时|天|日|月|年)"
    )
    expected = sorted(set(fact_pattern.findall(source_text)))
    if not expected:
        raise AssertionError("源中未识别到钱或时间事实")
    missing = [fact for fact in expected if fact not in content]
    if missing:
        raise AssertionError(f"候选缺少源中的钱或时间事实：{missing[:5]}")


def xlsx_2026_plan(
    case: dict,
    candidate: Path,
    *,
    source_resolver: SourceResolver = resolve_sources,
) -> None:
    rows = read_candidate_csv(candidate)
    sheets = _xlsx_sheets(source_resolver(case)[0][0])
    source_rows = sheets.get("大师建设计划进展(副本)")
    if not source_rows:
        raise AssertionError("源中未找到大师建设计划进展表")

    def select_plans(table_rows: list[list[str]]) -> list[tuple[str, str]]:
        header_index = _find_header(table_rows, "AI大师", "2026")
        header = [normalized(cell) for cell in table_rows[header_index]]
        master_index = next(
            i for i, cell in enumerate(header) if "AI大师" in cell
        )
        year_index = next(i for i, cell in enumerate(header) if "2026" in cell)
        return [
            (normalized(row[master_index]), normalized(row[year_index]))
            for row in table_rows[header_index + 1 :]
            if max(master_index, year_index) < len(row)
            and normalized(row[master_index])
            and normalized(row[year_index])
        ]

    expected = select_plans(source_rows)
    actual = select_plans(rows)
    if Counter(actual) != Counter(expected):
        raise AssertionError(
            f"2026 建设计划与源不一致（期望 {len(expected)} 项，实际 {len(actual)} 项）"
        )


def pdf_modules_no_audit(case: dict, candidate: Path) -> None:
    rows = read_candidate_csv(candidate)
    data = _nonempty_rows(rows)
    if len(data) < 3:
        raise AssertionError(f"功能模块行数异常（期望 ≥3，实际 {len(data)}）")
    # 禁止项：候选不得包含文档历史/审签信息
    for row in data:
        joined = normalized(" ".join(row))
        for forbidden in ("历史", "审签", "校对", "批准"):
            if forbidden in joined:
                raise AssertionError(f"候选包含禁止项「{forbidden}」：{row}")


ASSERTS = {
    "pdf_workload_detail": pdf_workload_detail,
    "pdf_workload_summary": pdf_workload_summary,
    "pdf_workload_detail_second_source": pdf_workload_detail_second_source,
    "pdf_workload_merge_three": pdf_workload_merge_three,
    "pdf_page3_table": pdf_page3_table,
    "pdf_first_document": pdf_first_document,
    "pdf_all_documents": pdf_all_documents,
    "pdf_function_modules": pdf_function_modules,
    "pdf_plan_topics": pdf_plan_topics,
    "pdf_miniapp_quiz_steps": pdf_miniapp_quiz_steps,
    "docx_core_indicators": docx_core_indicators,
    "docx_po_modules": docx_po_modules,
    "docx_region_elements": docx_region_elements,
    "docx_agents_scenarios": docx_agents_scenarios,
    "docx_quality_indicators": docx_quality_indicators,
    "docx_otd_schedule": docx_otd_schedule,
    "xlsx_sheet2_filter": xlsx_sheet2_filter,
    "xlsx_scene_planning": xlsx_scene_planning,
    "xlsx_plan2026_new": xlsx_plan2026_new,
    "xlsx_ktm_owner": xlsx_ktm_owner,
    "csv_sum_by_dept": csv_sum_by_dept,
    "csv_total_amount": csv_total_amount,
    "csv_filter_value": csv_filter_value,
    "csv_garbled_header_sum": csv_garbled_header_sum,
    "csv_merge_by_name": csv_merge_by_name,
    "csv_combined_total": csv_combined_total,
    "docx_two_docs_modules": docx_two_docs_modules,
    "csv_three_merge": csv_three_merge,
    "docx_money_time_extract": docx_money_time_extract,
    "xlsx_2026_plan": xlsx_2026_plan,
    "pdf_modules_no_audit": pdf_modules_no_audit,
}

# 只有这些规则会逐值对照源内容或执行精确计算；其余历史规则仍只可用于
# diagnostic 回归，不得因“结构非空/行数足够”进入正式 G1。
FORMAL_ASSERT_RULES = {
    "docx_agents_scenarios",
    "xlsx_sheet2_filter",
    "xlsx_scene_planning",
    "xlsx_plan2026_new",
    "csv_sum_by_dept",
    "csv_filter_value",
    "csv_garbled_header_sum",
    "docx_money_time_extract",
    "xlsx_2026_plan",
}


def run_assert(
    case: dict,
    candidate_path: Path,
    *,
    source_resolver: SourceResolver = resolve_sources,
) -> None:
    rule = case["assert_rule"]
    if rule not in ASSERTS:
        raise AssertionError(f"未知断言规则：{rule}")
    source_aware_asserts = {
        "docx_po_modules": docx_po_modules,
        "docx_agents_scenarios": docx_agents_scenarios,
        "xlsx_scene_planning": xlsx_scene_planning,
        "xlsx_plan2026_new": xlsx_plan2026_new,
        "docx_money_time_extract": docx_money_time_extract,
        "xlsx_2026_plan": xlsx_2026_plan,
    }
    if rule in source_aware_asserts:
        source_aware_asserts[rule](
            case,
            candidate_path,
            source_resolver=source_resolver,
        )
        return
    ASSERTS[rule](case, candidate_path)


def validate_all_assertions_importable() -> None:
    """启动时自检：fixtures 中所有 assert_rule 都有实现。"""
    fixtures_path = EVALS_ROOT / "fixtures.json"
    fixtures = json.loads(fixtures_path.read_text(encoding="utf-8"))
    missing = [
        c["id"]
        for c in fixtures["cases"]
        if c["assert_rule"] not in ASSERTS
    ]
    if missing:
        raise AssertionError(f"未实现的断言规则：{missing}")


def formal_assertion_gaps(fixtures: dict) -> tuple[str, ...]:
    """阻止弱历史规则进入正式盲保留集。"""

    return tuple(
        f"{case['id']} 使用非正式强度断言：{case['assert_rule']}"
        for case in fixtures.get("cases", ())
        if case.get("assert_rule") not in FORMAL_ASSERT_RULES
    )
