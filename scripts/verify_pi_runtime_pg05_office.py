# -*- coding: utf-8 -*-
"""用真实 Word/Excel 上传对象验证 PG-05 的跨格式泛化能力。"""
from __future__ import annotations

import argparse
import asyncio
from collections import Counter
import csv
from dataclasses import dataclass
import hashlib
import json
from pathlib import Path
import re
import secrets
import sys
import tempfile
import unicodedata

from docx import Document
from openpyxl import load_workbook

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from src.agentic_runtime.models import (
    PermissionProfile,
    PiRuntimeRequest,
    RuntimeEvent,
    SourceInput,
    VerificationStatus,
)
from src.agentic_runtime.pi_runtime import PiRuntime, PiRuntimeError
from src.config.settings import settings


MAX_TIMEOUT_SECONDS = 7200


@dataclass(frozen=True)
class OfficeValidationBatch:
    case: str
    batch_id: str
    owner_id: str
    upload_id: str
    repeat: int
    timeout_seconds: int

    def task_id(self, run_number: int) -> str:
        # 批次身份必须进入任务身份，重复或并发执行不能复用另一批的运行资源。
        return f"pg05_{self.case}_{self.batch_id}_{run_number}"

    def allocate_execution_root(
        self,
        base_root: Path,
        run_number: int,
    ) -> Path:
        batch_root = (
            base_root / f"pi-runtime-pg05-{self.case}" / self.batch_id
        )
        batch_root.mkdir(parents=True, exist_ok=True)
        return Path(
            tempfile.mkdtemp(prefix=f"run-{run_number}-", dir=batch_root)
        )


def _positive_timeout(value: str) -> int:
    try:
        timeout_seconds = int(value)
    except ValueError as exc:
        raise argparse.ArgumentTypeError(
            "timeout-seconds 必须是正整数"
        ) from exc
    if timeout_seconds <= 0:
        raise argparse.ArgumentTypeError("timeout-seconds 必须是正整数")
    # 验收任务允许长时间运行，但仍需硬上限，防止配置错误形成无界资源占用。
    if timeout_seconds > MAX_TIMEOUT_SECONDS:
        raise argparse.ArgumentTypeError(
            f"timeout-seconds 不能超过 {MAX_TIMEOUT_SECONDS}"
        )
    return timeout_seconds


def _safe_identifier(field_name: str):
    def parse(value: str) -> str:
        if not value or not value.strip():
            raise argparse.ArgumentTypeError(f"{field_name} 不能为空")
        if value != value.strip():
            raise argparse.ArgumentTypeError(
                f"{field_name} 不能包含首尾空白"
            )
        if "/" in value or "\\" in value or value in {".", ".."}:
            raise argparse.ArgumentTypeError(
                f"{field_name} 不能包含路径分隔符"
            )
        return value

    return parse


def _resolve_upload(owner_id: str, upload_id: str) -> tuple[Path, dict]:
    upload_root = Path(settings.data_prep_upload_root)
    if not upload_root.is_absolute():
        upload_root = PROJECT_ROOT / upload_root
    upload_root = upload_root.resolve()
    owner_objects = (upload_root / owner_id / "objects").resolve()
    # CLI 必须绑定明确 Owner；即使目录被替换为链接也不能逃出上传根目录。
    if not owner_objects.is_relative_to(upload_root):
        raise PermissionError("指定 Owner 的上传目录无效")
    meta_path = owner_objects / f"{upload_id}.meta"
    if not meta_path.is_file():
        raise ValueError("指定 Owner 下未找到上传对象")
    metadata = json.loads(meta_path.read_text(encoding="utf-8"))
    for field_name in ("upload_id", "user_id", "original_name", "sha256"):
        value = metadata.get(field_name)
        if not isinstance(value, str) or not value.strip():
            raise ValueError(
                f"上传对象元数据字段 {field_name} 不能为空"
            )
    if str(metadata.get("user_id") or "") != owner_id:
        raise PermissionError("上传对象元数据与指定 Owner 不一致")
    if str(metadata.get("upload_id") or "") != upload_id:
        raise ValueError("上传对象元数据与指定 upload_id 不一致")
    source = meta_path.with_suffix("").resolve()
    if source.parent != owner_objects:
        raise PermissionError("上传对象路径超出指定 Owner 目录")
    if not source.is_file():
        raise ValueError("上传对象源文件不存在")
    actual = hashlib.sha256(source.read_bytes()).hexdigest()
    if actual != metadata["sha256"]:
        raise ValueError("真实上传对象与元数据哈希不一致")
    return source, metadata


def _normalized(value: object) -> str:
    return re.sub(
        r"\s+",
        "",
        unicodedata.normalize("NFKC", str(value or "")),
    )


def _word_source_text(source: Path) -> str:
    document = Document(source)
    parts = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            parts.extend(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )
    return "\n".join(parts)


def _assert_word_summary(candidate: Path, source: Path) -> None:
    content = candidate.read_text(encoding="utf-8-sig")
    source_text = _word_source_text(source)
    required_categories = (
        "费用",
        "服务",
        "验收",
        "保密",
        "知识产权",
        "违约",
    )
    missing = [
        category
        for category in required_categories
        if category not in content
    ]
    if missing:
        raise AssertionError(f"TXT 缺少商务条款类别：{missing}")
    if len(content.strip()) < 800:
        raise AssertionError("TXT 过短，未形成可用的商务条款汇总")
    # 这是防止“把整份 Word 原封不动转成 TXT”的独立验收门，不影响摘要写法。
    if len(content) >= len(source_text) * 0.75:
        raise AssertionError("TXT 接近全文长度，疑似整份文档照搬")


def _expected_excel_rows(source: Path) -> Counter[tuple[str, str]]:
    with source.open("rb") as handle:
        workbook = load_workbook(handle, read_only=True, data_only=True)
        try:
            worksheet = workbook["工作表2"]
            rows = Counter()
            for row in worksheet.iter_rows(min_row=2, values_only=True):
                if _normalized(row[5] if len(row) > 5 else "") != "1月":
                    continue
                scene = _normalized(row[4] if len(row) > 4 else "")
                maturity = _normalized(row[6] if len(row) > 6 else "")
                if scene:
                    rows[(scene, maturity)] += 1
            return rows
        finally:
            workbook.close()


def _candidate_excel_rows(candidate: Path) -> Counter[tuple[str, str]]:
    with candidate.open(encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.reader(handle))
    if not rows:
        raise AssertionError("CSV 是空文件")
    header_index = next(
        (
            index
            for index, row in enumerate(rows)
            if any(
                "应用场景" in _normalized(cell)
                and "三级" in _normalized(cell)
                for cell in row
            )
            and any("目标成熟度" in _normalized(cell) for cell in row)
        ),
        None,
    )
    if header_index is None:
        raise AssertionError("CSV 缺少约定的两列表头")
    header = [_normalized(cell) for cell in rows[header_index]]
    scene_index = next(
        index
        for index, cell in enumerate(header)
        if "应用场景" in cell and "三级" in cell
    )
    maturity_index = next(
        index
        for index, cell in enumerate(header)
        if "目标成熟度" in cell
    )
    actual: Counter[tuple[str, str]] = Counter()
    for row in rows[header_index + 1 :]:
        scene = _normalized(
            row[scene_index] if scene_index < len(row) else ""
        )
        maturity = _normalized(
            row[maturity_index] if maturity_index < len(row) else ""
        )
        if scene:
            actual[(scene, maturity)] += 1
    return actual


def _assert_excel_filter(candidate: Path, source: Path) -> None:
    expected = _expected_excel_rows(source)
    actual = _candidate_excel_rows(candidate)
    if actual != expected:
        missing = list((expected - actual).elements())[:5]
        extra = list((actual - expected).elements())[:5]
        raise AssertionError(
            "CSV 与真实工作表筛选结果不一致："
            f"期望 {sum(expected.values())} 行，实际 {sum(actual.values())} 行；"
            f"缺少 {missing}；多出 {extra}"
        )


def _case_contract(case: str) -> tuple[str, str]:
    if case == "word":
        return (
            "按费用、服务、验收、保密、知识产权、违约六类，提取并汇总"
            "这个 Word 里的商务条款，输出一个 TXT；不要复制整份文档。",
            "txt",
        )
    return (
        "只读取 Excel 的【工作表2】，筛选【计划交付月份】为【1月】的记录，"
        "只输出【应用场景（三级）】和【目标成熟度】两列，生成一个 CSV；"
        "其他工作表、其他月份和说明文字都不要。",
        "csv",
    )


async def _run_once(
    *,
    batch: OfficeValidationBatch,
    source: Path,
    metadata: dict,
    run_number: int,
) -> None:
    objective, output_format = _case_contract(batch.case)
    execution_root = batch.allocate_execution_root(
        PROJECT_ROOT / ".pytest-tmp",
        run_number,
    )
    request = PiRuntimeRequest(
        user_id=batch.owner_id,
        task_id=batch.task_id(run_number),
        revision=1,
        objective_text=objective,
        requested_output_formats=(output_format,),
        sources=(
            SourceInput(
                upload_id=str(metadata["upload_id"]),
                original_name=str(metadata["original_name"]),
                host_path=source,
                sha256=str(metadata["sha256"]),
                media_type=str(
                    metadata.get("media_type")
                    or "application/octet-stream"
                ),
            ),
        ),
        permission_profile=PermissionProfile.STANDARD,
        model=settings.llm_model_name,
        base_url=settings.llm_base_url,
        api_key="local-runtime",
    )

    async def event_sink(event: RuntimeEvent) -> None:
        print(
            f"[batch:{batch.batch_id}][{batch.case}:{run_number}] "
            f"{event.event_type}: {event.summary}",
            flush=True,
        )

    result = await PiRuntime(
        execution_root=execution_root,
        timeout_seconds=batch.timeout_seconds,
    ).start(request, on_event=event_sink)
    if len(result.candidates) != 1:
        raise AssertionError(
            f"必须只生成一个 {output_format.upper()}，"
            f"实际 {len(result.candidates)} 个候选"
        )
    candidate = result.candidates[0]
    if candidate.format != output_format:
        raise AssertionError(
            f"候选格式不是 {output_format.upper()}：{candidate.format}"
        )
    if batch.case == "word":
        _assert_word_summary(candidate.host_path, source)
    else:
        _assert_excel_filter(candidate.host_path, source)
    if (
        result.verification is None
        or result.verification.status is not VerificationStatus.PASSED
    ):
        raise AssertionError(
            "独立验证未通过："
            + (
                result.verification.model_dump_json()
                if result.verification
                else "无验证报告"
            )
        )
    print(
        f"[batch:{batch.batch_id}][{batch.case}:{run_number}] "
        f"PASS: {candidate.host_path}",
        flush=True,
    )


async def _main(batch: OfficeValidationBatch) -> None:
    print(
        f"[batch:{batch.batch_id}] START case={batch.case} "
        f"owner={batch.owner_id} repeat={batch.repeat} "
        f"timeout_seconds={batch.timeout_seconds}",
        flush=True,
    )
    source, metadata = _resolve_upload(batch.owner_id, batch.upload_id)
    for run_number in range(1, batch.repeat + 1):
        await _run_once(
            batch=batch,
            source=source,
            metadata=metadata,
            run_number=run_number,
        )


def _failure_category(exc: Exception) -> str:
    if isinstance(exc, PermissionError):
        return "permission_denied"
    if isinstance(exc, (ValueError, FileNotFoundError, KeyError)):
        return "input_invalid"
    if isinstance(exc, AssertionError):
        return "verification_failed"
    if isinstance(exc, PiRuntimeError):
        cause: BaseException | None = exc
        while cause is not None:
            if isinstance(cause, TimeoutError):
                return "runtime_timeout"
            cause = cause.__cause__
        return "runtime_failed"
    return "unexpected_error"


def _redact_host_paths(message: str) -> str:
    # 验证日志可以保留端点与错误语义，但不得暴露 Owner 文件的宿主绝对路径。
    return re.sub(
        r"(?<![A-Za-z0-9])(?:[A-Za-z]:[\\/]|\\\\)[^\r\n；;]+",
        "<host-path>",
        message,
    )


def main() -> int:
    parser = argparse.ArgumentParser(
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument("--case", required=True, choices=("word", "excel"))
    parser.add_argument(
        "--owner-id",
        required=True,
        type=_safe_identifier("owner-id"),
    )
    parser.add_argument(
        "--upload-id",
        required=True,
        type=_safe_identifier("upload-id"),
    )
    parser.add_argument("--repeat", type=int, default=1, choices=range(1, 4))
    parser.add_argument(
        "--timeout-seconds",
        type=_positive_timeout,
        default=1800,
        help="每次 Pi 执行允许的最长秒数",
    )
    args = parser.parse_args()
    batch = OfficeValidationBatch(
        case=args.case,
        batch_id=secrets.token_hex(8),
        owner_id=args.owner_id,
        upload_id=args.upload_id,
        repeat=args.repeat,
        timeout_seconds=args.timeout_seconds,
    )
    try:
        asyncio.run(_main(batch))
    except Exception as exc:
        print(
            f"[batch:{batch.batch_id}] FAILED "
            f"category={_failure_category(exc)} "
            f"error={type(exc).__name__}: "
            f"{_redact_host_paths(str(exc))}",
            file=sys.stderr,
            flush=True,
        )
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
