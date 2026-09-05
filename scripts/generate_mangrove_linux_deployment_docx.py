from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Mangrove_Linux服务器完整部署方案.docx"
LOGO = ROOT / "logos" / "公司LOGO.png"

GREEN = "147D64"
DARK = "20352F"
LIGHT_GREEN = "EAF5F1"
LIGHT_GRAY = "F4F6F7"
MID_GRAY = "66736F"
RED = "B42318"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D5DEDB", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_run_font(run, name: str = "微软雅黑", size: float | None = None, bold: bool | None = None,
                 color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MID_GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=MID_GRAY)


def add_toc(document: Document) -> None:
    p = document.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "在 Word 中右键此处并选择“更新域”，即可生成目录。"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, placeholder, fld_end])


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True


def add_para(document: Document, text: str = "", bold_prefix: str | None = None) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=DARK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.25
        set_run_font(p.add_run(item))


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.25
        set_run_font(p.add_run(item))


def add_code(document: Document, code: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F5F4")
    p_pr.append(shd)
    for i, line in enumerate(code.strip("\n").splitlines()):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=8.5, color="24332E")


def add_note(document: Document, title: str, text: str, warning: bool = False) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF1E8" if warning else LIGHT_GREEN)
    set_cell_border(cell, "E7B38E" if warning else "9BCBBC", "6")
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}：")
    set_run_font(r, bold=True, color=RED if warning else GREEN)
    set_run_font(p.add_run(text))
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, GREEN)
        set_cell_border(cell, "FFFFFF", "4")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), size=9, bold=True, color="FFFFFF")
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(values):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, "FFFFFF" if row_idx % 2 == 0 else LIGHT_GRAY)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(str(text)), size=8.8)
            if widths:
                cell.width = Cm(widths[i])
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    for level, size, color in ((1, 18, GREEN), (2, 14, DARK), (3, 11.5, GREEN)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6)
    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(10.5)


def build_document() -> None:
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.1)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    # Cover
    for _ in range(2):
        doc.add_paragraph()
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Cm(5.8))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(32)
    set_run_font(p.add_run("Mangrove（红树林）"), size=28, bold=True, color=GREEN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Linux 服务器完整部署方案"), size=24, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    set_run_font(p.add_run("生产架构 · 安装配置 · 数据迁移 · 安全运维 · 升级回滚"), size=11.5, color=MID_GRAY)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("南京华苏科技"), size=12, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(f"版本：1.0    日期：{date.today().isoformat()}"), size=10, color=MID_GRAY)
    doc.add_page_break()

    add_heading(doc, "文档说明", 1)
    add_para(doc, "本方案基于 Mangrove v1.1.0 当前代码、配置和依赖进行部署审计，目标是在单台 Linux 服务器上形成可运行、可维护、可备份、可升级回滚的生产环境。")
    add_note(doc, "重要结论", "仓库根目录现有 Dockerfile 和 docker-compose.yml 仍属于旧版 Browser Agent/Streamlit，不能直接作为当前 Mangrove 的生产部署文件。首期推荐采用“systemd 运行 Mangrove 主进程 + Docker 运行辅助采集服务 + Nginx 提供 HTTPS”的混合架构。", True)
    add_table(doc, ["文档项", "内容"], [
        ["适用版本", "Mangrove v1.1.0 及相近版本"],
        ["推荐系统", "Ubuntu Server 24.04 LTS，x86_64"],
        ["部署模式", "单机、单 FastAPI 进程"],
        ["主要受众", "系统管理员、实施工程师、运维与项目负责人"],
        ["不包含", "多节点高可用、Kubernetes、分布式任务队列的具体实现"],
    ], [3.5, 12])

    add_heading(doc, "目录", 1)
    add_toc(doc)
    doc.add_page_break()

    add_heading(doc, "1. 推荐生产架构", 1)
    add_para(doc, "推荐架构将公网入口、核心应用、运行数据和可选采集服务分层，既保持现有代码的单实例语义，也便于后续逐步容器化。")
    add_code(doc, """
用户浏览器
    │ HTTPS 443
    ▼
Nginx
    │ 127.0.0.1:8088
    ▼
Mangrove FastAPI（systemd、单进程）
    ├── SQLite：用户 / 会话 / 调度 / 检查点
    ├── 持久化目录：模板 / 教训 / 报告 / 日志
    ├── SearXNG :8080       ┐
    ├── Firecrawl :3002     ├── Docker 辅助服务
    ├── RSSHub :1200        ┘
    ├── MediaCrawler 独立 Python venv
    └── 外部 LLM / Embedding / Reranker
""")
    add_bullets(doc, [
        "公网只访问 Nginx 的 80/443，内部服务不直接暴露。",
        "FastAPI 同源托管构建后的 React 前端，生产环境不运行 Vite 5173。",
        "调度器、Cookie 巡检器、模板巡检器均随 FastAPI lifespan 启动。",
        "当前 SQLite、后台任务登记和调度器均为单实例模型，不能简单增加 Uvicorn worker。",
    ])

    add_heading(doc, "2. 服务器与网络准备", 1)
    add_table(doc, ["使用场景", "CPU", "内存", "磁盘"], [
        ["基础功能，不启用 Firecrawl/浏览器", "4 核", "8 GB", "50 GB SSD"],
        ["完整单机部署", "8 核", "16 GB", "100 GB SSD"],
        ["Firecrawl、MediaCrawler 高频使用", "16 核", "32 GB", "200 GB SSD"],
    ], [7.3, 2.3, 2.5, 3.2])
    add_bullets(doc, [
        "Ubuntu Server 24.04 LTS，建议 x86_64。",
        "Python 3.13；Node.js 20 或以上仅用于构建前端。",
        "Docker Engine 和 Docker Compose Plugin。",
        "准备域名、DNS 解析和 HTTPS 证书。",
        "服务器可访问所使用的 LLM、搜索、邮件和 Slack 服务。",
        "本地大模型推荐部署在独立推理节点；不在本机运行模型时无需 GPU。",
    ])
    add_table(doc, ["端口", "用途", "公网策略"], [
        ["22", "SSH", "仅允许管理 IP"],
        ["80/443", "Nginx HTTP/HTTPS", "允许公网访问"],
        ["8088", "Mangrove FastAPI", "仅监听 127.0.0.1"],
        ["8080", "SearXNG", "仅监听 127.0.0.1"],
        ["1200", "RSSHub", "仅监听 127.0.0.1"],
        ["3002", "Firecrawl API", "仅监听 127.0.0.1"],
        ["5432/6379/5672", "Firecrawl PostgreSQL/Redis/RabbitMQ", "不映射公网"],
    ], [2.4, 6.7, 6.2])
    add_note(doc, "防火墙提醒", "Docker 发布端口可能绕过 UFW 普通规则。生产环境除 UFW/firewalld 外，还应检查 iptables 的 DOCKER-USER 链。", True)

    add_heading(doc, "3. 部署前置检查", 1)
    add_heading(doc, "3.1 冻结干净版本", 2)
    add_para(doc, "部署前应确认现有 scheduler、API、前端任务页、模板和教训数据改动，提交后打不可变部署标签。服务器只部署确定的 tag 或 commit。")
    add_code(doc, """
git status
git tag -a v1.1.0-prod.1 -m "Mangrove Linux production release"
git push origin v1.1.0-prod.1
""")
    add_heading(doc, "3.2 配置与许可检查", 2)
    add_bullets(doc, [
        "重新生成 Linux 环境变量，禁止直接照搬 Windows 盘符路径。",
        "检查 webui.db 中 runtime_config 是否覆盖 .env。",
        "所有在聊天、日志或历史文件中出现过的密钥均应轮换。",
        "Firecrawl 采用 AGPL-3.0；对外提供 SaaS 前应完成许可证评估。",
        "MediaCrawler 当前说明为非商业学习用途；商业化前应取得授权或替换数据源。",
    ])

    add_heading(doc, "4. Linux 目录规划", 1)
    add_code(doc, """
/opt/mangrove/
├── releases/
│   ├── v1.1.0-prod.1/
│   └── v1.1.0-prod.2/
├── current -> releases/v1.1.0-prod.1
├── venv/
└── mediacrawler-venv/

/srv/mangrove/
├── data/
├── downloads/
├── logs/
├── memory/
├── mediacrawler-data/
└── backups/

/etc/mangrove/
└── mangrove.env
""")
    add_code(doc, """
sudo useradd --system --create-home \
  --home-dir /opt/mangrove \
  --shell /usr/sbin/nologin mangrove

sudo mkdir -p /opt/mangrove/releases
sudo mkdir -p /srv/mangrove/{data,downloads,logs,memory,mediacrawler-data,backups}
sudo mkdir -p /etc/mangrove
sudo chown -R mangrove:mangrove /opt/mangrove /srv/mangrove
sudo chmod 750 /etc/mangrove
""")

    add_heading(doc, "5. 安装 Mangrove", 1)
    add_heading(doc, "5.1 获取代码", 2)
    add_code(doc, """
sudo -u mangrove git clone \
  --branch v1.1.0-prod.1 \
  <Git仓库地址> \
  /opt/mangrove/releases/v1.1.0-prod.1

sudo -u mangrove ln -s \
  /opt/mangrove/releases/v1.1.0-prod.1 \
  /opt/mangrove/current

cd /opt/mangrove/current
sudo -u mangrove git submodule update --init --recursive
""")
    add_heading(doc, "5.2 安装 Python 依赖", 2)
    add_code(doc, """
sudo -u mangrove python3.13 -m venv /opt/mangrove/venv
sudo -u mangrove /opt/mangrove/venv/bin/python -m pip install \
  --upgrade pip setuptools wheel
sudo -u mangrove /opt/mangrove/venv/bin/pip install \
  -r /opt/mangrove/current/requirements.txt

sudo /opt/mangrove/venv/bin/playwright install-deps chromium
sudo -u mangrove /opt/mangrove/venv/bin/playwright install chromium
sudo -u mangrove /opt/mangrove/venv/bin/scrapling install
""")
    add_para(doc, "若暂时不使用 Camoufox 隐身浏览器，可跳过 scrapling install，并在配置中设置 SCRAPLING_STEALTH_ENABLED=False。")
    add_heading(doc, "5.3 构建 React 前端", 2)
    add_code(doc, """
cd /opt/mangrove/current/frontend
sudo -u mangrove npm ci
sudo -u mangrove npm run build

test -f /opt/mangrove/current/frontend/dist/index.html
""")
    add_note(doc, "生产规则", "生产环境不运行 npm run dev。FastAPI 会同源托管 frontend/dist，用户只访问 8088 或 Nginx 的 HTTPS 地址。")

    add_heading(doc, "6. 数据持久化设计", 1)
    add_table(doc, ["路径", "用途", "重要性"], [
        ["data/webui.db", "用户、会话、反馈、个人记忆、运行时配置", "必须"],
        ["data/scheduler.db", "定时任务与执行历史", "必须"],
        ["data/checkpoints.sqlite", "LangGraph 断点续跑状态", "必须"],
        ["data/app.db", "用户确认后的业务入库结果", "按需"],
        ["data/templates/", "自学习分析模板", "必须"],
        ["data/lessons/", "失败教训库", "必须"],
        ["downloads/", "报告、JSON、trace", "必须"],
        ["memory/", "全局偏好", "必须"],
        ["logs/", "运行日志", "建议"],
        ["MediaCrawler data/", "社媒采集运行数据", "按需"],
    ], [4.6, 8, 2.6])
    add_code(doc, """
sudo -u mangrove rsync -a /opt/mangrove/current/data/ /srv/mangrove/data/
sudo -u mangrove rsync -a /opt/mangrove/current/memory/ /srv/mangrove/memory/
""")
    add_para(doc, "随后将发布目录中的 data、downloads、logs、memory 指向 /srv/mangrove 下对应目录。每次发布新版本时重新建立软链接，并用 rsync --ignore-existing 合并新版种子模板和教训。")
    add_note(doc, "SQLite 限制", "SQLite 数据库必须放在服务器本地 SSD，不建议放在 NFS、SMB 等网络文件系统。当前架构保持单进程、单节点。", True)

    add_heading(doc, "7. 生产环境变量", 1)
    add_para(doc, "创建 /etc/mangrove/mangrove.env。以下模板不包含真实密钥，部署时应按实际环境替换。")
    add_code(doc, """
# 模型
LLM_DEFAULT_PROVIDER=deepseek
DEEPSEEK_API_KEY=<真实密钥>
DEEPSEEK_BASE_URL=https://api.deepseek.com/v1
DEEPSEEK_MODEL=deepseek-chat

# 本地模型（可选）
LLM_MODEL_NAME=<模型名>
LLM_BASE_URL=http://<模型服务器IP>:<端口>/v1
LLM_API_KEY=local
LOCAL_ENABLE_THINKING=True
LOCAL_MAX_TOKENS=32768

# 主服务
API_HOST=127.0.0.1
API_PORT=8088
LOG_LEVEL=INFO
TZ=Asia/Shanghai

# 持久化
WEBUI_DB_PATH=/srv/mangrove/data/webui.db
SCHEDULER_DB_PATH=/srv/mangrove/data/scheduler.db
CHECKPOINT_DB_PATH=/srv/mangrove/data/checkpoints.sqlite
DATABASE_URL=sqlite:////srv/mangrove/data/app.db
SCHEDULER_ENABLED=True
CHECKPOINT_ENABLED=True

# 安全
JWT_SECRET=<至少64位随机字符串>
JWT_EXPIRE_HOURS=168
WEBUI_ALLOW_REGISTER=False
WEBUI_CORS_ORIGINS=https://mangrove.example.com

# 辅助采集服务
SEARXNG_BASE_URL=http://127.0.0.1:8080
FIRECRAWL_BASE_URL=http://127.0.0.1:3002
FIRECRAWL_API_KEY=
RSSHUB_BASE_URL=http://127.0.0.1:1200

# MediaCrawler
MEDIACRAWLER_PATH=/opt/mangrove/current/external/MediaCrawler/repo
MEDIACRAWLER_PYTHON=/opt/mangrove/mediacrawler-venv/bin/python
MC_ENABLE_CDP_MODE=False

# 搜索增强（可选）
TAVILY_API_KEY=
ANYSEARCH_API_KEY=

# 语义召回（可选）
EMBEDDING_ENABLED=False
EMBEDDING_BASE_URL=
EMBEDDING_API_KEY=
EMBEDDING_MODEL=
RERANK_BASE_URL=
RERANK_API_KEY=
RERANK_MODEL=
""")
    add_code(doc, """
openssl rand -hex 48
sudo chown root:mangrove /etc/mangrove/mangrove.env
sudo chmod 640 /etc/mangrove/mangrove.env
""")
    add_note(doc, "runtime_config", "管理员在配置中心保存的全局值存放在 data/webui.db，并在启动时覆盖 .env。从 Windows 迁移后必须检查旧路径、旧模型名和旧端口。", True)

    add_heading(doc, "8. 辅助采集服务部署", 1)
    add_heading(doc, "8.1 SearXNG", 2)
    add_code(doc, """
cd /opt/mangrove/current/docker/searxng
sudo docker compose up -d
""")
    add_para(doc, "将端口映射调整为 127.0.0.1:8080:8080，禁止直接暴露公网。")
    add_heading(doc, "8.2 RSSHub", 2)
    add_code(doc, """
sudo docker run -d \
  --name mangrove-rsshub \
  --restart unless-stopped \
  -p 127.0.0.1:1200:1200 \
  diygod/rsshub:<固定版本>
""")
    add_heading(doc, "8.3 Firecrawl", 2)
    add_code(doc, """
cd /opt/mangrove/current/external/firecrawl
sudo docker compose pull api playwright-service nuq-postgres redis rabbitmq
sudo docker compose up -d api nuq-postgres
""")
    add_bullets(doc, [
        "将 API 端口限制为 127.0.0.1:3002。",
        "给 PostgreSQL、RabbitMQ、Redis 添加命名持久卷。",
        "固定所有镜像版本或 digest，避免生产环境长期使用 latest。",
        "ALLOW_LOCAL_WEBHOOKS=true 会放宽 SSRF 防护，只能在 Firecrawl 不对公网开放时启用。",
    ])
    add_para(doc, "16 GB 内存服务器建议将 Firecrawl 并发调整为：")
    add_code(doc, """
NUM_WORKERS_PER_QUEUE=2
CRAWL_CONCURRENT_REQUESTS=2
MAX_CONCURRENT_JOBS=2
BROWSER_POOL_SIZE=2
""")
    add_heading(doc, "8.4 MediaCrawler", 2)
    add_code(doc, """
sudo -u mangrove python3.13 -m venv /opt/mangrove/mediacrawler-venv
sudo -u mangrove /opt/mangrove/mediacrawler-venv/bin/pip install \
  -r /opt/mangrove/current/external/MediaCrawler/repo/requirements.txt
sudo /opt/mangrove/mediacrawler-venv/bin/playwright install-deps chromium
sudo -u mangrove /opt/mangrove/mediacrawler-venv/bin/playwright install chromium
""")
    add_para(doc, "Linux 无人值守环境建议先关闭 CDP 模式，使用 MediaCrawler 自带的 Playwright 浏览器。若必须复用真实 Chrome，应将主应用部署在宿主机并单独设计浏览器进程和用户数据目录。")

    add_heading(doc, "9. systemd 服务", 1)
    add_para(doc, "创建 /etc/systemd/system/mangrove.service：")
    add_code(doc, """
[Unit]
Description=Mangrove Data Collection Agent
After=network-online.target docker.service
Wants=network-online.target
Requires=docker.service

[Service]
Type=simple
User=mangrove
Group=mangrove
WorkingDirectory=/opt/mangrove/current
EnvironmentFile=/etc/mangrove/mangrove.env
Environment=PYTHONUNBUFFERED=1
Environment=TZ=Asia/Shanghai
ExecStart=/opt/mangrove/venv/bin/python -m src.api.main
Restart=always
RestartSec=5
TimeoutStopSec=120
KillSignal=SIGTERM
LimitNOFILE=65535
UMask=0077
NoNewPrivileges=true
PrivateTmp=true

[Install]
WantedBy=multi-user.target
""")
    add_code(doc, """
sudo systemctl daemon-reload
sudo systemctl enable --now mangrove
sudo systemctl status mangrove
sudo journalctl -u mangrove -f
""")
    add_note(doc, "单实例要求", "不要使用 --reload、多个 Uvicorn worker 或多个 Mangrove 副本，否则可能出现定时任务重复执行、后台任务状态不一致和 SQLite 竞争。", True)

    add_heading(doc, "10. Nginx 与 HTTPS", 1)
    add_code(doc, """
server {
    listen 80;
    server_name mangrove.example.com;
    location / { return 301 https://$host$request_uri; }
}

server {
    listen 443 ssl http2;
    server_name mangrove.example.com;

    ssl_certificate     /etc/letsencrypt/live/mangrove.example.com/fullchain.pem;
    ssl_certificate_key /etc/letsencrypt/live/mangrove.example.com/privkey.pem;
    client_max_body_size 20m;

    location /assets/ {
        proxy_pass http://127.0.0.1:8088;
        proxy_set_header Host $host;
        expires 7d;
    }

    location / {
        proxy_pass http://127.0.0.1:8088;
        proxy_http_version 1.1;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
        proxy_buffering off;
        proxy_cache off;
        proxy_read_timeout 3600s;
        proxy_send_timeout 3600s;
    }
}
""")
    add_para(doc, "Mangrove 使用 SSE 流式输出。Nginx 必须关闭响应缓冲并延长读取超时，否则前端可能无法实时显示节点事件。")

    add_heading(doc, "11. 首次管理员初始化", 1)
    add_numbered(doc, [
        "暂时不要开放 Nginx 公网入口。",
        "完成 WebUI 数据库显式迁移，并配置至少 32 字节随机 JWT_SECRET。",
        "维护者在服务器交互终端运行：python -m src.api.bootstrap_admin --database data/webui.db --username maintainer。",
        "按提示隐藏输入并确认至少 12 位密码；已有超级管理员时跳过，不覆盖账号。",
        "保持 WEBUI_ALLOW_REGISTER=False，通过 SSH 隧道登录管理员验证。",
        "完成管理员初始化后再开放 HTTPS 域名。",
    ])
    add_note(doc, "初始化边界", "首次管理员只由本机维护者显式建立。公开注册即使开启也只创建待审批普通用户。", True)

    add_heading(doc, "12. Windows 数据迁移", 1)
    add_para(doc, "迁移前停止 Windows Mangrove，确保 SQLite 不再写入。需要复制：")
    add_bullets(doc, [
        "data/webui.db、data/scheduler.db、data/checkpoints.sqlite、data/app.db（如存在）",
        "data/templates/、data/lessons/",
        "downloads/",
        "memory/user-preferences.md",
    ])
    add_note(doc, "数据库识别", "当前实际用户数据库是 data/webui.db；项目根目录的小型 webui.db 更可能是历史遗留文件，不应默认作为生产库迁移。")
    add_code(doc, """
sudo chown -R mangrove:mangrove /srv/mangrove
sudo find /srv/mangrove -type d -exec chmod 750 {} \;
sudo find /srv/mangrove -type f -exec chmod 640 {} \;
sudo timedatectl set-timezone Asia/Shanghai
timedatectl
""")
    add_para(doc, "迁移后检查 runtime_config、MediaCrawler 路径、本地模型地址、CORS 域名、定时任务时区和 Cookie 有效性。")

    add_heading(doc, "13. 上线验证清单", 1)
    add_heading(doc, "13.1 健康检查", 2)
    add_code(doc, """
curl http://127.0.0.1:8088/api/health
curl http://127.0.0.1:8080
curl http://127.0.0.1:1200
curl http://127.0.0.1:3002
""")
    add_heading(doc, "13.2 功能验收", 2)
    add_table(doc, ["序号", "验证项", "预期结果"], [
        ["1", "HTTPS 页面与登录", "页面正常、证书可信、登录成功"],
        ["2", "模型目录", "供应商和模型列表可加载"],
        ["3", "普通对话", "LLM 正常返回"],
        ["4", "URL 抓取", "生成 Markdown、JSON 和 trace"],
        ["5", "关键词搜索", "SearXNG 或其他后端生效"],
        ["6", "定时任务", "按服务器时区触发并留下历史"],
        ["7", "下载", "报告和 JSON 可下载且权限隔离"],
        ["8", "服务重启", "会话、调度和配置仍存在"],
        ["9", "断点续跑", "中断任务可从检查点恢复"],
        ["10", "Cookie 校验", "按平台返回有效/失效/无法判断"],
        ["11", "SSE", "节点事件持续、实时显示"],
        ["12", "管理功能", "反馈、模板、教训、用户管理正常"],
    ], [1.4, 5.3, 9.1])

    add_heading(doc, "14. 备份与恢复", 1)
    add_para(doc, "每天备份 SQLite、模板、教训、个人/全局记忆、报告、加密后的环境变量及 Firecrawl 数据。最稳妥的方式是短暂停机冷备：")
    add_code(doc, """
sudo systemctl stop mangrove
sudo tar -C /srv -czf \
  /srv/mangrove/backups/mangrove-$(date +%F-%H%M).tar.gz \
  mangrove/data mangrove/downloads mangrove/memory
sudo systemctl start mangrove
""")
    add_para(doc, "若不能停机，应使用 sqlite3 的 .backup 命令逐库生成一致性备份，不要直接复制正在使用的 .db、-wal、-shm 文件。Firecrawl PostgreSQL 使用 pg_dump 或卷快照。")
    add_table(doc, ["周期", "保留建议"], [
        ["每日", "保留最近 7 天"],
        ["每周", "保留最近 4 周"],
        ["每月", "保留最近 6 个月"],
        ["异地备份", "至少保留一份加密副本"],
        ["恢复演练", "至少每季度一次"],
    ], [4, 11])

    add_heading(doc, "15. 升级与回滚", 1)
    add_heading(doc, "15.1 升级步骤", 2)
    add_numbered(doc, [
        "备份 SQLite 数据库和所有持久化目录。",
        "将新 tag 拉取到 /opt/mangrove/releases/<version>。",
        "安装或更新 Python 依赖。",
        "执行 npm ci 和 npm run build。",
        "合并新版本种子模板、教训和技能。",
        "停止 Mangrove，切换 current 软链接。",
        "启动服务并执行健康检查和功能冒烟。",
    ])
    add_heading(doc, "15.2 回滚步骤", 2)
    add_numbered(doc, [
        "停止 Mangrove。",
        "如果数据库结构已变化，恢复升级前数据库备份。",
        "将 current 指回旧版本。",
        "重新启动并验证登录、采集、调度和下载。",
    ])
    add_note(doc, "回滚原则", "不要只回滚代码而保留可能不兼容的新数据库。每次升级都必须先备份。", True)

    add_heading(doc, "16. 后续全容器化路线", 1)
    add_para(doc, "功能冻结后可制作新的生产 Dockerfile 和 Compose。目标要求如下：")
    add_bullets(doc, [
        "使用多阶段构建生成 React frontend/dist。",
        "使用 Python 3.13 运行镜像，启动命令为 python -m src.api.main。",
        "安装 Chromium、Playwright、中文字体、Scrapling/Camoufox。",
        "暴露 8088，不再使用旧版 8000/8501。",
        "为 data、downloads、logs、memory 和 MediaCrawler data 挂载持久卷。",
        "设置 init: true、足够的 shm_size、健康检查和日志轮转。",
        "保持单 worker；辅助服务走内部 Docker 网络。",
        "固定所有镜像版本或 digest。",
        "仅由 Nginx/Caddy 暴露 80/443。",
        "MediaCrawler CDP 模式默认关闭。",
    ])
    add_para(doc, "多节点高可用不能只靠复制容器实现。需要将 webui、scheduler、checkpoint 等状态迁移到共享数据库，引入 Redis/任务队列、分布式锁和明确的任务所有权。")

    add_heading(doc, "17. 常见故障排查", 1)
    add_table(doc, ["现象", "重点检查"], [
        ["8088 无监听", "journalctl、Python import、依赖安装、端口占用、冷启动时间"],
        ["页面仍是旧版本", "frontend/dist 是否重新 npm run build"],
        ["修改 .env 不生效", "webui.db 的 runtime_config 是否覆盖"],
        ["SSE 最后一次性显示", "Nginx proxy_buffering 是否关闭"],
        ["定时任务重复执行", "是否启动多个进程或多个副本"],
        ["Cookie 验证失败", "Cookie 过期、平台风控、浏览器依赖、服务器出口 IP"],
        ["本地模型 502", "LAN 路由、代理、模型地址、trust_env 分流"],
        ["模型返回空内容", "思考模型 max_tokens 是否过小"],
        ["SQLite locked", "是否多 worker、数据是否位于网络盘、备份方式是否错误"],
        ["Firecrawl 内存不足", "降低 workers、并发、浏览器池；增加 swap 不能代替内存"],
    ], [5.7, 9.4])

    add_heading(doc, "18. 参考资料", 1)
    refs = [
        "项目 README_AGENT.md 与 AGENTS.md",
        "项目 .env.example、src/api/main.py、src/config/settings.py",
        "Docker Engine for Ubuntu：https://docs.docker.com/engine/install/ubuntu/",
        "Docker Compose Plugin：https://docs.docker.com/compose/install/linux/",
        "Docker Volumes：https://docs.docker.com/engine/storage/volumes/",
        "Nginx proxy module：https://nginx.org/en/docs/http/ngx_http_proxy_module.html",
    ]
    add_bullets(doc, refs)

    # Headers/footers after content so all sections receive them.
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run_font(header.add_run("Mangrove · Linux 服务器完整部署方案"), size=8.5, color=MID_GRAY)
        add_page_number(section.footer.paragraphs[0])

    # Enable update fields on open, allowing Word to refresh TOC.
    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
